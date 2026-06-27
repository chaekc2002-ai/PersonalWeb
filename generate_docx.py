from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row):
            row_cells[i].text = text

def generate_supplier_checklist(filename):
    doc = Document()
    
    title = doc.add_heading('학습지원 소프트웨어 필수기준 체크리스트(공급자용)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('□ 제품/서비스 개요', level=1)
    overview_data = [
        ['제품/서비스명', '학급용 웹앱 통합 포털', '공급자', '채관철'],
        ['접속경로', 'https://(할당된 배포 주소)', '', ''],
        ['주요 내용 및 기능·특장점', '◦ 학생 개인정보 수집 없는 안전한 링크 허브\n◦ 별도 로그인/설치 없이 브라우저에서 바로 사용 가능', '', '']
    ]
    create_table(doc, ['항목', '내용', '항목2', '내용2'], overview_data)
    
    doc.add_paragraph()
    doc.add_heading('□ 개인정보보호 기준 충족여부', level=1)
    headers = ['선정기준', '세부 내용', '확인 (■충족/□미충족/□해당없음)', '증빙']
    data = [
        ['1. 최소처리 원칙 준수', '1-1. 개인정보가 최소한으로 수집되는가?', '■충족', '(개인정보 처리방침 제2조 등 - 개인정보 수집 없음)'],
        ['1. 최소처리 원칙 준수', '1-2. 개인정보 수집·이용 목적이 기재되어 있는가?', '■충족', '(개인정보 처리방침 제1조 등)'],
        ['1. 최소처리 원칙 준수', '1-3. 개인정보 수집항목, 보유기간 등이 기재되어 있는가?', '■충족', '(개인정보 처리방침 제2, 3조 등)'],
        ['2. 개인정보 안전조치 의무', '2-1. 개인정보 안전성 확보에 필요한 조치 사항이 기재되어 있는가?', '■충족', '(개인정보 처리방침 제6조 등)'],
        ['3. 열람/정정/삭제 절차', '3-1. 이용자에게 언제든지 자신의 정보를 열람·정정·삭제·처리정지를 요구할 수 있는 절차가 안내되어 있는가?', '■충족', '(개인정보 처리방침 제7조 등)'],
        ['4. 만14세 미만 아동 보호', '4-1. 만 14세 미만 아동의 경우 법정대리인 동의 등 아동의 개인정보 보호를 위한 절차가 마련되어 있는가?', '■해당없음', '(정보 수집 없음)'],
        ['5. 보호책임자/제3자제공 등', '5-1. 개인정보 보호책임자 관련 정보가 안내되어 있는가?', '■충족', '(개인정보 처리방침 제8조 등)'],
        ['5. 보호책임자/제3자제공 등', '5-2. 개인정보 제3자 제공에 관한 정보가 기재되어 있는가?', '■해당없음', '(정보 제공 없음)'],
        ['5. 보호책임자/제3자제공 등', '5-3. 개인정보 위·수탁관계에 관한 정보가 기재되어 있는가?', '■해당없음', '(정보 위탁 없음)']
    ]
    create_table(doc, headers, data)
    
    doc.add_paragraph()
    doc.add_heading('□ 작성일 및 문의처', level=1)
    doc.add_paragraph('작성일: 2026. 06. 27.')
    doc.add_paragraph('문의처: chaekc2002@sen.go.kr')
    
    doc.save(filename)

def generate_school_checklist(filename):
    doc = Document()
    
    title = doc.add_heading('학습지원 소프트웨어 선정기준 체크리스트 (학교용)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('□ 필수기준 (개인정보보호)', level=1)
    headers = ['선정기준', '세부 내용', '확인', '증빙자료']
    data = [
        ['1. 최소처리 원칙 준수', '1-1. 개인정보가 최소한으로 수집되는가?', '■충족', '개인정보를 일절 수집하지 않고 익명 접근을 보장합니다.'],
        ['1. 최소처리 원칙 준수', '1-2. 개인정보 수집·이용 목적이 기재되어 있는가?', '■충족', '수집 및 이용 목적(수집 않음)을 명확히 기재하였습니다.'],
        ['1. 최소처리 원칙 준수', '1-3. 개인정보 수집항목, 보유기간 등이 기재되어 있는가?', '■충족', '보유 기간이 존재하지 않음을 상세히 안내하고 있습니다.'],
        ['2. 안전조치 의무', '2-1. 개인정보 안전성 확보에 필요한 조치 사항이 기재되어 있는가?', '■충족', '비밀번호 접근 통제 등 기술적 보호 조치 사항을 포함하고 있습니다.'],
        ['3. 열람/정정/삭제 절차', '3-1. 이용자에게 언제든지 정보를 열람·삭제 요구할 수 있는 절차가 안내되어 있는가?', '■충족', '수집된 정보가 없음을 안내하여 권리 행사 절차를 갈음하고 있습니다.'],
        ['4. 만14세 미만 아동 보호', '4-1. 아동의 개인정보 보호를 위한 절차가 마련되어 있는가?', '■충족', '정보 수집 자체가 없어 별도의 동의 절차가 불필요함을 안내합니다.'],
        ['5. 보호책임자 명시 등', '5-1. 개인정보 보호책임자 관련 정보가 안내되어 있는가?', '■충족', '개인정보 보호책임자(채관철)를 명시하고 있습니다.'],
        ['5. 보호책임자 명시 등', '5-2. 개인정보 제3자 제공에 관한 정보가 기재되어 있는가?', '■해당없음', '개인정보를 제3자에게 제공하지 않습니다.'],
        ['5. 보호책임자 명시 등', '5-3. 개인정보 위·수탁관계에 관한 정보가 기재되어 있는가?', '■해당없음', '위탁 처리하지 않음을 기재하고 있습니다.']
    ]
    create_table(doc, headers, data)
    
    doc.add_paragraph()
    doc.add_heading('□ 선택기준', level=1)
    data_optional = [
        ['1. 교육목표 적합성', '1-1. 수업 목표와 학생의 학습 수준에 적합한 내용과 기능을 제공하는가?', '■충족', 'AI 윤리 가이드 확약 및 학습 도구 통합 접근 기능 제공'],
        ['2. 콘텐츠 안전성', '2-1. 학습 콘텐츠가 정확하고 신뢰할 수 있으며, 안전한가?', '■충족', '교사가 검증한 안전한 링크만을 제공합니다.'],
        ['3. 사용 환경 적합성', '3-1. 학교의 기기·네트워크 환경에서 안정적으로 사용할 수 있는가?', '■충족', '웹 표준 기반으로 별도 앱 설치 없이 작동합니다.'],
        ['4. 접근성 및 사용성', '4-1. 교사와 학생이 필요한 기능과 자료에 쉽게 접근하고 활용할 수 있는가?', '■충족', '직관적인 UI와 원클릭 이동을 지원합니다.']
    ]
    create_table(doc, headers, data_optional)

    doc.save(filename)

if __name__ == '__main__':
    generate_supplier_checklist('에듀집 탑재용 체크리스트(공급자용).docx')
    generate_school_checklist('학교용 필수기준체크리스트.docx')
